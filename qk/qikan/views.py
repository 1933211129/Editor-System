from django.shortcuts import render
from django.http import JsonResponse, HttpResponseNotModified
from django.views.decorators.http import require_http_methods
import requests
from bs4 import BeautifulSoup
import re
import json
import hashlib
import time
import random
from .models import (
    Mail, JournalPeriodOne, JournalPeriodTwo, JournalPeriodThree,
    JournalPeriodFour, JournalPeriodFive, Invoice, ArticleSchedule, PeriodMapping, Contact, Reviewer, Todo, Recipient, RecipientGroup, RecipientGroupMembership, UpdateLog
)
from django.db.models import Q, Count
from django.contrib.auth import authenticate, login
from django.contrib.auth.models import User
from django.views.decorators.csrf import csrf_exempt
from django.utils.dateparse import parse_date
import os
from django.conf import settings
from django.utils import timezone
from docx import Document
from django.http import FileResponse
from docxtpl import DocxTemplate
from .generation_demo import generate_notice_pdfs
from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from .serializers import ContactSerializer, ReviewerSerializer
from .mail_sender import EmailSender
from datetime import datetime
from email.utils import parsedate_to_datetime

# 固定发信配置（如需修改请更新以下常量）
# 注意：将真实的邮箱与授权码填入下方常量。当前占位值用于运行，但无法实际发送成功。
SMTP_SERVER_FIXED = 'mail.cstnet.cn'
SMTP_PORT_FIXED = 465
SENDER_EMAIL_FIXED = 'jcip@iscas.ac.cn'
SENDER_PASSWORD_FIXED = 'GtVkDjHbcH8^@sFc'

# 附件临时存放目录（不入库，纯临时文件）
TEMP_UPLOAD_DIR = os.path.join(getattr(settings, 'BASE_DIR', os.getcwd()), 'temp_email_uploads')
os.makedirs(TEMP_UPLOAD_DIR, exist_ok=True)

# ========== 附件上传相关 API ==========
# ========== 更新日志 API ==========
@csrf_exempt
@require_http_methods(["GET"])
def get_latest_update_log(request):
    """返回最新一条更新日志（内容与时间）。若无数据返回空。"""
    try:
        latest = UpdateLog.objects.order_by('-updated_at').first()
        if not latest:
            return JsonResponse({
                'status': 'success',
                'data': None
            })
        return JsonResponse({
            'status': 'success',
            'data': {
                'content': latest.content,
                'updatedAt': latest.updated_at.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def upload_email_attachments(request):
    """
    POST /api/email/upload/
    表单字段: files (支持多文件)
    返回: { status, files: [{name, size}] }
    """
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': '不支持的请求方法'}, status=405)
    try:
        files = request.FILES.getlist('files')
        saved = []
        for f in files:
            # 基于原始文件名保存，如冲突则追加序号
            base = os.path.basename(f.name)
            name, ext = os.path.splitext(base)
            candidate = base
            idx = 1
            while os.path.exists(os.path.join(TEMP_UPLOAD_DIR, candidate)):
                candidate = f"{name}({idx}){ext}"
                idx += 1
            dest_path = os.path.join(TEMP_UPLOAD_DIR, candidate)
            with open(dest_path, 'wb') as out:
                for chunk in f.chunks():
                    out.write(chunk)
            saved.append({'name': candidate, 'size': os.path.getsize(dest_path)})
        return JsonResponse({'status': 'success', 'files': saved})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def list_email_uploads(request):
    """
    GET /api/email/uploads/
    列出临时上传目录下的文件
    """
    if request.method != 'GET':
        return JsonResponse({'status': 'error', 'message': '不支持的请求方法'}, status=405)
    try:
        files = []
        for fname in os.listdir(TEMP_UPLOAD_DIR):
            fpath = os.path.join(TEMP_UPLOAD_DIR, fname)
            if os.path.isfile(fpath):
                files.append({'name': fname, 'size': os.path.getsize(fpath)})
        return JsonResponse({'status': 'success', 'files': files})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def clear_email_uploads(request):
    """
    POST /api/email/uploads/clear/
    清空临时上传目录（谨慎使用）
    """
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': '不支持的请求方法'}, status=405)
    try:
        removed = 0
        for fname in os.listdir(TEMP_UPLOAD_DIR):
            fpath = os.path.join(TEMP_UPLOAD_DIR, fname)
            if os.path.isfile(fpath):
                os.remove(fpath)
                removed += 1
        return JsonResponse({'status': 'success', 'removed': removed})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

# Create your views here.
def process_journal(url="http://jcip.cipsc.org.cn/CN/home", api_url="http://43.138.77.63/crawl_qk"):
    """
    调用 Flask API 爬取期刊信息。
    
    参数:
        url (str): 目标期刊页面的 URL，默认为 "http://jcip.cipsc.org.cn/CN/home"。
        api_url (str): Flask API 的地址，默认为 "http://43.138.77.63/crawl_qk"。
        
    返回:
        dict: 从 API 返回的爬取结果。
    """
    try:
        # 构造请求参数
        params = {"url": url}
        
        # 发送 GET 请求
        response = requests.get(api_url, params=params)
        
        # 检查响应状态码
        if response.status_code == 200:
            result = response.json()
            if result["status"] == "success":
                return result["data"]
            else:
                print(f"API 调用失败: {result['message']}")
        else:
            print(f"HTTP 请求失败，状态码: {response.status_code}")
    except Exception as e:
        print(f"发生错误: {e}")


def generate_markdown_content(data):
    """
    动态生成 Markdown 内容的函数。
    
    参数:
        data (dict): 嵌套字典数据，包含 H1 和对应的 H2 列表。
    
    返回:
        str: 生成的 HTML 格式内容。
    """
    # 颜色列表
    color_list = [
        '#8577c1', '#61c1eb', '#feb20f', '#9d3c3f', '#fd7faa', '#009933', '#663300',
        '#f2c96d', '#4a69bd', '#9a8c98', '#78e08f', '#ff6348', '#58d89d', '#f78fb3',
        '#b3b6b7', '#f5f0c0', '#7f8fa6', '#a27b5c', '#8fbc8f', '#f0e130'
    ]
    content = ""
    
    # 遍历每个 H1 及其对应的 H2 列表
    for title_h1, h2_list in data.items():
        # 添加 H1 标题
        content += f"""
<h1 style="
    font-family: 'Arial', sans-serif; /* 使用楷体 */
    font-size: 1.65em;
    color: #000; /* 黑色文字 */
    text-align: center;
    margin: 20px auto; /* 居中对齐 */
    display: block; /* 设置为块级元素 */
    padding: 1px 60px; /* 这里可以调整高度和宽度 */
    background: #fff; /* 白色背景 */
    box-shadow:
        4px 4px 8px rgba(0, 0, 0, 0.2), /* 阴影效果 */
        inset 0 0 10px rgba(0, 0, 0, 0.1); /* 内部阴影 */
    border-radius: 5px; /* 圆角 */
">
  {title_h1}
</h1>

{'<br><br><br>'}  <!-- H1 和第一个 H2 之间空出四个换行 -->
"""

        # 遍历 H2 列表
        for i, h2_data in enumerate(h2_list):
            # 随机选择一个颜色作为二级标题的颜色
            h2_color = random.choice(color_list)
            
            # 提取 H2 数据
            title_h2 = h2_data["title_cn"]
            authors = h2_data["author_cn"]
            abstract = h2_data["abstract_cn"]
            keywords = ", ".join(h2_data["keywords_cn"])
            citation_zh = h2_data["reference_cn"]
            citation_en = h2_data["reference_en"]
            full_text_link = h2_data["url"]

            # 添加 H2 标题及内容
            content += f"""
<h2 style="font-family: 'Arial', sans-serif; font-size: 1.4em; color: {h2_color}; text-shadow: 1px 1px 3px rgba(0, 0, 0, 0.2); text-align: left; margin-top: 20px; margin-bottom: 10px;">
  ✦ {title_h2}
</h2>

<p style="font-family: 'Arial', sans-serif; font-size: 1em; color: #555; line-height: 1.6; margin-top: 10px;">
  <strong>作&nbsp;&nbsp;&nbsp;&nbsp;者:</strong> {authors}
</p>

<p style="font-family: 'Arial', sans-serif; font-size: 1em; color: #555; line-height: 1.6; margin-top: 10px;">
  <strong>摘&nbsp;&nbsp;&nbsp;&nbsp;要:</strong> {abstract}
</p>

<p style="font-family: 'Arial', sans-serif; font-size: 1em; color: #555; line-height: 1.6; margin-top: 10px;">
  <strong>关键词 :</strong> {keywords}
</p>

<p style="font-family: 'Arial', sans-serif; font-size: 1em; color: #555; line-height: 1.6; margin-top: 10px;">
  <strong>引用格式 :</strong> {citation_zh}
</p>

<p style="font-family: 'Arial', sans-serif; font-size: 1em; color: #555; line-height: 1.6; margin-top: 10px;">
  <strong></strong> {citation_en}
</p>

<p style="font-family: 'Arial', sans-serif; font-size: 1em; color: #555; line-height: 1.6; margin-top: 10px;">
  <strong>全文链接：</strong>
  <a href="{full_text_link}" style="color: #007BFF; text-decoration: none;">点击下载</a>
</p>
"""

            # 如果不是最后一个 H2，则在 H2 之间空出两个换行
            if i < len(h2_list) - 1:
                content += '<br><br>'

        # H2 和下一个 H1 之间空出四个换行
        content += '<br><br><br>'

    print("生成完成")
    return content

@require_http_methods(["GET"])
def get_journal_content(request):
    """
    API endpoint to get journal content
    """
    try:
        # 获取期刊数据
        journal_data = process_journal()
        
        # 生成 Markdown 内容
        markdown_content = generate_markdown_content(journal_data)
        
        # 返回成功响应
        return JsonResponse({
            'status': 'success',
            'data': {
                'markdown': markdown_content,
                'raw_data': journal_data
            }
        })
    except Exception as e:
        # 返回错误响应
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

def get_email_history(request):
    search_query = request.GET.get('search', '')
    
    # 基础查询（注意 Mail.date 为字符串，且格式混合，不能直接可靠按字典序排序）
    emails_qs = Mail.objects.all()
    
    # 如果有搜索关键词，添加搜索条件
    if search_query:
        emails_qs = emails_qs.filter(
            Q(sender__icontains=search_query) |
            Q(recipient__icontains=search_query) |
            Q(subject__icontains=search_query) |
            Q(body__icontains=search_query) |
            Q(attachments__icontains=search_query)
        )
    
    def _parse_mail_date_str(s: str) -> datetime:
        """尽量把字符串日期解析为有时区的 datetime。失败则返回 1970-01-01 UTC，用于稳定排序。"""
        if not s:
            return datetime(1970, 1, 1, tzinfo=timezone.utc)
        # 1) ISO 格式
        try:
            ds = s.replace('Z', '+00:00')
            dt = datetime.fromisoformat(ds)
            if dt.tzinfo is None:
                dt = timezone.make_aware(dt)
            return dt
        except Exception:
            pass
        # 2) RFC2822 格式（如 Wed, 13 Sep 2023 11:06:37 +0800）
        try:
            dt = parsedate_to_datetime(s)
            if dt is not None and dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt or datetime(1970, 1, 1, tzinfo=timezone.utc)
        except Exception:
            pass
        return datetime(1970, 1, 1, tzinfo=timezone.utc)

    # 抓取对象后在 Python 里按解析后的时间倒序排序
    emails_objs = list(emails_qs)
    # 预解析成 (obj, dt) 并用时间戳排序，保证稳定
    parsed = [(e, _parse_mail_date_str(e.date)) for e in emails_objs]
    parsed.sort(key=lambda t: t[1].timestamp() if t[1] else 0, reverse=True)
    
    # 如果没有搜索，只返回最新的前50条
    if not search_query:
        parsed = parsed[:50]
    
    email_list = []
    for email, dt in parsed:
        email_list.append({
            'type': email.type,
            'sender': email.sender,
            'recipient': email.recipient,
            'subject': email.subject,
            'date': email.date,
            'timestamp': dt.isoformat() if dt else '',
            'body': email.body,
            'attachments': email.attachments
        })
    
    return JsonResponse(email_list, safe=False)

@csrf_exempt
def user_register(request):
    print("收到注册请求")  # 添加调试日志
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            print("注册数据:", data)  # 添加调试日志
            username = data.get('username')
            email = data.get('email')  # 新增 email 字段
            password = data.get('password')
            
            # 检查必填字段
            if not username or not password or not email:
                return JsonResponse({
                    'status': 'error',
                    'message': '用户名、邮箱和密码不能为空'
                }, status=400)
            
            # 检查用户名是否已存在
            if User.objects.filter(username=username).exists():
                return JsonResponse({
                    'status': 'error',
                    'message': '用户名已存在'
                }, status=400)
                
            # 检查邮箱是否已存在
            if User.objects.filter(email=email).exists():
                return JsonResponse({
                    'status': 'error',
                    'message': '邮箱已被注册'
                }, status=400)
            
            # 创建用户
            try:
                user = User.objects.create_user(
                    username=username,
                    email=email,
                    password=password
                )
                return JsonResponse({
                    'status': 'success',
                    'message': '注册成功'
                })
            except Exception as e:
                return JsonResponse({
                    'status': 'error',
                    'message': f'创建用户失败: {str(e)}'
                }, status=500)
                
        except json.JSONDecodeError:
            return JsonResponse({
                'status': 'error',
                'message': '无效的请求数据格式'
            }, status=400)
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'服务器错误: {str(e)}'
            }, status=500)
    
    return JsonResponse({
        'status': 'error',
        'message': '不支持的请求方法'
    }, status=405)

@csrf_exempt
def user_login(request):
    print("收到登录请求")  # 添加调试日志
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            print("登录数据:", data)  # 添加调试日志
            username = data.get('username')  # 这里接收的是邮箱
            password = data.get('password')
            
            if not username or not password:
                return JsonResponse({
                    'status': 'error',
                    'message': '邮箱和密码不能为空'
                }, status=400)
            
            # 先尝试用邮箱查找用户
            try:
                user = User.objects.get(email=username)
                username = user.username  # 获取真实用户名
            except User.DoesNotExist:
                pass  # 如果找不到，继续使用输入的用户名
            
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)  # 添加这行来更新last_login
                return JsonResponse({
                    'status': 'success',
                    'token': 'dummy_token',
                    'username': user.username,
                    'message': '登录成功'
                })
            else:
                return JsonResponse({
                    'status': 'error',
                    'message': '邮箱或密码错误'
                }, status=401)
                
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': f'服务器错误: {str(e)}'
            }, status=500)

    return JsonResponse({
        'status': 'error',
        'message': '不支持的请求方法'
    }, status=405)

# 获取所有数据
@csrf_exempt
def get_journal_data(request, period_number):
    if request.method in ['GET', 'POST']:
        try:
            # 根据期数选择对应的模型
            model_map = {
                1: JournalPeriodOne,
                2: JournalPeriodTwo,
                3: JournalPeriodThree,
                4: JournalPeriodFour,
                5: JournalPeriodFive
            }
            
            JournalModel = model_map.get(period_number)
            if not JournalModel:
                return JsonResponse({'status': 'error', 'message': '无效的期数'}, status=400)
            
            # 如果是 POST 请求，处理导入
            if request.method == 'POST':
                new_records = []

                # 优先处理文件名列表（避免上传文件内容）
                try:
                    if not request.FILES:
                        payload = json.loads(request.body or '{}')
                        filenames = payload.get('filenames', [])
                        if isinstance(filenames, list) and filenames:
                            for name in filenames:
                                journal = JournalModel.objects.create(
                                    filename=str(name),
                                    editor_in_charge='',
                                    page_fee='待更新',
                                    proof_status='待更新',
                                    first_second_proof_editor='',
                                    third_proof_editor='',
                                    final_proof_editor='',
                                    remarks=''
                                )
                                new_records.append({
                                    'id': journal.id,
                                    'title': journal.filename,
                                    'responsible': journal.editor_in_charge,
                                    'remarks': journal.remarks,
                                    'stages': {
                                        'fee': journal.page_fee,
                                        'proof': journal.proof_status
                                    },
                                    'editors': {
                                        'proof12': journal.first_second_proof_editor,
                                        'proof3': journal.third_proof_editor,
                                        'proofFinal': journal.final_proof_editor
                                    },
                                    'proofDates': {
                                        'editor': journal.editor_time or '',
                                        'proof': journal.proof_time or ''
                                    }
                                })
                except Exception:
                    # 忽略 JSON 解析失败，回退到文件上传方式
                    pass

                # 兼容旧逻辑：如果包含文件则按文件创建
                if not new_records:
                    files = request.FILES.getlist('files')
                    for file in files:
                        journal = JournalModel.objects.create(
                            filename=file.name,
                            editor_in_charge='',
                            page_fee='待更新',
                            proof_status='待更新',
                            first_second_proof_editor='',
                            third_proof_editor='',
                            final_proof_editor='',
                            remarks=''
                        )
                        new_records.append({
                            'id': journal.id,
                            'title': journal.filename,
                            'responsible': journal.editor_in_charge,
                            'remarks': journal.remarks,
                            'stages': {
                                'fee': journal.page_fee,
                                'proof': journal.proof_status
                            },
                            'editors': {
                                'proof12': journal.first_second_proof_editor,
                                'proof3': journal.third_proof_editor,
                                'proofFinal': journal.final_proof_editor
                            },
                            'proofDates': {
                                'editor': journal.editor_time or '',
                                'proof': journal.proof_time or ''
                            }
                        })

                return JsonResponse({'status': 'success', 'data': new_records})
                
            # GET 请求返回所有数据
            journals = JournalModel.objects.all()
            data = []
            for journal in journals:
                item = {
                    'id': journal.id,
                    'title': journal.filename,
                    'responsible': journal.editor_in_charge,
                    'remarks': journal.remarks,
                    'stages': {
                        'fee': journal.page_fee,
                        'proof': journal.proof_status
                    },
                    'editors': {
                        'proof12': journal.first_second_proof_editor,
                        'proof3': journal.third_proof_editor,
                        'proofFinal': journal.final_proof_editor
                    },
                    'proofDates': {
                        'editor': journal.editor_time,
                        'proof': journal.proof_time
                    }
                }
                data.append(item)
            return JsonResponse(data, safe=False)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': '不支持的请求方法'}, status=405)

@csrf_exempt
def send_bulk_email(request):
    """
    POST /api/email/send-bulk/
    请求体(JSON):
    {
      "default_subject": str,                 // 批量默认主题，可为空
      "subjects": [str],                      // 可选，与收件人等长，单独覆盖主题
      "html_body": str,                       // HTML 正文
      "receiver_emails": [str],
      "unique_attachment_names": [[str]],     // 每个收件人的附件文件名列表（已上传至服务器）
      "common_attachment_names": [str]        // 通用附件文件名列表（已上传至服务器）
    }
    说明：SMTP 发信配置已在后端写死（见模块常量），前端无需再传入/切换邮箱配置。
    返回: 每个收件人的发送结果。
    """
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': '不支持的请求方法'}, status=405)

    try:
        payload = json.loads(request.body or '{}')

        default_subject = payload.get('default_subject') or ''
        subjects = payload.get('subjects') or []
        html_body = payload.get('html_body') or ''
        receiver_emails = payload.get('receiver_emails') or []
        unique_attachment_names = payload.get('unique_attachment_names') or []
        common_attachment_names = payload.get('common_attachment_names') or []

        # 参数校验
        if not isinstance(receiver_emails, list) or not isinstance(unique_attachment_names, list):
            return JsonResponse({'status': 'error', 'message': 'receiver_emails 与 unique_attachments 必须为列表'}, status=400)
        if len(receiver_emails) != len(unique_attachment_names):
            return JsonResponse({'status': 'error', 'message': '收件人数量与个性化附件列数量不一致'}, status=400)

        mailer = EmailSender(
            smtp_server=SMTP_SERVER_FIXED,
            smtp_port=SMTP_PORT_FIXED,
            sender_email=SENDER_EMAIL_FIXED,
            sender_password=SENDER_PASSWORD_FIXED,
        )

        # 执行批量发送，收集结果
        results = []
        # 记录此次请求中引用到的临时文件（仅限本次 payload 指定的文件名）
        used_file_names = set()
        try:
            # 将待用文件名提前汇总，发送完成后统一清理
            for n in (common_attachment_names or []):
                if n:
                    used_file_names.add(os.path.basename(str(n)))
            for ua in (unique_attachment_names or []):
                if isinstance(ua, list):
                    for n in ua:
                        if n:
                            used_file_names.add(os.path.basename(str(n)))
        except Exception:
            # 聚合过程中出错不影响后续发送
            pass
        try:
            import smtplib
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText

            with smtplib.SMTP_SSL(mailer.smtp_server, mailer.smtp_port) as server:
                server.login(mailer.sender_email, mailer.sender_password)

                for i, receiver in enumerate(receiver_emails):
                    ok = True
                    error_msg = ''
                    attachment_status = []  # [{name, attached, error}]
                    try:
                        msg = MIMEMultipart()
                        msg['From'] = mailer.sender_email
                        msg['To'] = receiver
                        # 选择主题：优先 subjects[i]，否则默认主题
                        subject_i = ''
                        if isinstance(subjects, list) and i < len(subjects) and subjects[i]:
                            subject_i = str(subjects[i])
                        else:
                            subject_i = default_subject
                        msg['Subject'] = subject_i
                        msg.attach(MIMEText(html_body, 'html', 'utf-8'))

                        # 拼装附件: 通用 + 个性化（基于文件名到服务器路径的映射）
                        ua = unique_attachment_names[i] if i < len(unique_attachment_names) else []
                        if not isinstance(ua, list):
                            ua = []
                        # 将文件名映射为服务器路径
                        def to_path(name):
                            return os.path.join(TEMP_UPLOAD_DIR, os.path.basename(name))
                        common_paths = [to_path(n) for n in (common_attachment_names or [])]
                        unique_paths = [to_path(n) for n in (ua or [])]
                        all_paths = list(dict.fromkeys(list(common_paths) + list(unique_paths)))
                        for p in all_paths:
                            fname = os.path.basename(p)
                            try:
                                mailer._add_attachment_to_msg(msg, p)
                                attachment_status.append({'name': fname, 'attached': True, 'error': ''})
                            except Exception as e:
                                # 附件失败不阻断整封邮件
                                attachment_status.append({'name': fname, 'attached': False, 'error': str(e)})

                        server.sendmail(mailer.sender_email, receiver, msg.as_string())
                    except Exception as e:
                        ok = False
                        error_msg = str(e)
                    # 记录到邮件历史
                    try:
                        _save_mail_history(
                            mail_type='bulk',
                            sender=mailer.sender_email,
                            recipient=receiver,
                            subject=subject_i,
                            body=html_body,
                            attachments=attachment_status,
                            error=error_msg,
                            success=ok,
                        )
                    except Exception:
                        # 历史记录失败不影响响应
                        pass

                    results.append({'recipient': receiver, 'success': ok, 'error': error_msg, 'attachments': attachment_status})

        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'SMTP 发送阶段失败: {str(e)}'}, status=500)

        # 发送阶段结束后，尝试清理本次使用过的临时附件文件
        cleanup = {'removed': 0, 'failed': 0}
        for fname in used_file_names:
            try:
                fpath = os.path.join(TEMP_UPLOAD_DIR, fname)
                if os.path.isfile(fpath):
                    os.remove(fpath)
                    cleanup['removed'] += 1
            except Exception:
                cleanup['failed'] += 1

        return JsonResponse({'status': 'success', 'results': results, 'cleanup': cleanup})

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def _save_mail_history(mail_type: str, sender: str, recipient: str, subject: str, body: str, attachments, error: str, success: bool):
    """将单封发信记录写入 `Mail` 表。
    - attachments: 列表，元素可为字符串或 {name, attached, error}
    - 失败时会把错误前缀写入 body 便于检索（不改表结构）
    """
    # 将附件规范化为 JSON 可存对象
    norm_atts = []
    if isinstance(attachments, list):
        for a in attachments:
            if isinstance(a, dict):
                norm_atts.append({
                    'name': a.get('name', ''),
                    'attached': bool(a.get('attached', False)),
                    'error': a.get('error', ''),
                })
            else:
                norm_atts.append({'name': str(a), 'attached': True, 'error': ''})
    date_str = timezone.now().isoformat(timespec='seconds')
    body_to_save = body or ''
    if not success and error:
        body_to_save = f"[SendError] {error}\n\n" + body_to_save
    # 写入外部表（models.Mail 是 unmanaged=True）
    try:
        Mail.objects.create(
            type=mail_type,
            sender=sender,
            recipient=recipient,
            subject=subject or '',
            date=date_str,
            body=body_to_save,
            attachments=norm_atts,
        )
    except Exception:
        # 历史写入失败不影响主流程
        pass

@csrf_exempt
def get_invoice_data(request):
    try:
        invoices = Invoice.objects.all().order_by('-id')
        data = [{
            'id': invoice.id,
            'manuscriptId': invoice.article_id,
            'article': invoice.article,
            'amount': invoice.amount,
            'invoiceDate': invoice.payment_date.isoformat() if invoice.payment_date else '',
            'paymentMethod': invoice.payment_method,
            'type': invoice.type,
            'tag': invoice.tag,
            'company': invoice.company,
            'taxId': invoice.tax_id,
            'contact': invoice.contact,
            'email': invoice.email,
            'phone': invoice.phone,
            'notes': invoice.notes
        } for invoice in invoices]
        return JsonResponse(data, safe=False)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def update_invoice(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            record_id = data.get('id')
            updates = data.get('updates', {})
            
            print(f"Received updates for record {record_id}:", updates)  # 添加调试日志
            
            try:
                invoice = Invoice.objects.get(id=record_id)
                
                field_mapping = {
                    'manuscriptId': 'article_id',
                    'article': 'article',
                    'amount': 'amount',
                    'invoiceDate': 'payment_date',
                    'paymentMethod': 'payment_method',
                    'type': 'type',
                    'tag': 'tag',
                    'company': 'company',
                    'taxId': 'tax_id',
                    'contact': 'contact',
                    'email': 'email',
                    'phone': 'phone',
                    'notes': 'notes'
                }
                
                for field, value in updates.items():
                    if field in field_mapping:
                        db_field = field_mapping[field]
                        print(f"Updating field {field} -> {db_field} = {value}")  # 添加调试日志
                        
                        if db_field == 'payment_date' and value:
                            try:
                                value = parse_date(value)
                            except Exception as e:
                                print(f"Date parsing error: {str(e)}")
                                continue  # 跳过无效的日期，继续处理其他字段
                        
                        setattr(invoice, db_field, value)
                
                invoice.save()
                return JsonResponse({'status': 'success'})
                
            except Invoice.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': f'找不到ID为 {record_id} 的记录'
                }, status=404)
                
        except Exception as e:
            print(f"Error updating invoice: {str(e)}")  # 添加调试日志
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def create_invoice(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            invoice = Invoice.objects.create(
                article_id=data['manuscriptId'],
                article=data.get('article', ''),
                amount=data.get('amount', ''),
                payment_date=parse_date(data.get('invoiceDate', '')) if data.get('invoiceDate') else None,
                payment_method=data.get('paymentMethod', ''),
                type=data.get('type', ''),
                tag=data.get('tag', ''),
                company=data.get('company', ''),
                tax_id=data.get('taxId', ''),
                contact=data.get('contact', ''),
                email=data.get('email', ''),
                phone=data.get('phone', ''),
                notes=data.get('notes', '')
            )
            return JsonResponse({
                'status': 'success',
                'id': invoice.id,
                'manuscriptId': invoice.article_id
            })
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def delete_invoice(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            article_id = data.get('manuscriptId')
            Invoice.objects.get(article_id=article_id).delete()
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def clear_invoice_data(request):
    if request.method == 'POST':
        try:
            Invoice.objects.all().delete()
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def import_invoice_data(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body).get('data', [])
            for item in data:
                Invoice.objects.create(
                    article_id=item['manuscriptId'],
                    article=item.get('article', ''),
                    amount=item.get('amount', ''),
                    payment_date=parse_date(item.get('invoiceDate', '')) if item.get('invoiceDate') else None,
                    payment_method=item.get('paymentMethod', ''),
                    type=item.get('type', ''),
                    tag=item.get('tag', ''),
                    company=item.get('company', ''),
                    tax_id=item.get('taxId', ''),
                    contact=item.get('contact', ''),
                    email=item.get('email', ''),
                    phone=item.get('phone', ''),
                    notes=item.get('notes', '')
                )
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def get_article_schedule_data(request):
    try:
        schedules = ArticleSchedule.objects.all()
        data = [{
            'id': schedule.id,
            'filename': schedule.filename,
            'schedule': schedule.schedule,
            'confirmed': schedule.confirmed,
            'notes': schedule.notes
        } for schedule in schedules]
        return JsonResponse(data, safe=False)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def update_article_schedule(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            schedule_id = data.get('id')
            field = data.get('field')
            value = data.get('value')

            schedule = ArticleSchedule.objects.get(id=schedule_id)
            setattr(schedule, field, value)
            schedule.save()

            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def create_article_schedule(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            schedule = ArticleSchedule.objects.create(
                filename=data.get('filename', ''),
                schedule=data.get('schedule', ''),
                confirmed=data.get('confirmed', '待更新'),
                notes=data.get('notes', '')
            )
            return JsonResponse({
                'status': 'success',
                'id': schedule.id
            })
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def delete_article_schedule(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            schedule_id = data.get('id')
            ArticleSchedule.objects.get(id=schedule_id).delete()
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def clear_article_schedule_data(request):
    if request.method == 'POST':
        try:
            ArticleSchedule.objects.all().delete()
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def import_article_schedule_data(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body).get('data', [])
            for item in data:
                ArticleSchedule.objects.create(
                    filename=item.get('filename', ''),
                    schedule=item.get('schedule', ''),
                    confirmed=item.get('confirmed', '待更新'),
                    notes=item.get('notes', '')
                )
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def import_files(request, period):
    if request.method == 'POST':
        try:
            # 根据期数选择对应的模型
            model_map = {
                1: JournalPeriodOne,
                2: JournalPeriodTwo,
                3: JournalPeriodThree,
                4: JournalPeriodFour,
                5: JournalPeriodFive
            }
            
            JournalModel = model_map.get(period)
            if not JournalModel:
                return JsonResponse({'status': 'error', 'message': '无效的期数'}, status=400)

            files = request.FILES.getlist('files')
            new_records = []
            
            for file in files:
                # 创建新记录
                journal = JournalModel.objects.create(
                    filename=file.name,
                    editor_in_charge='',
                    page_fee='待更新',
                    proof_status='待更新',
                    remarks=''
                )
                
                # 构造返回数据
                new_records.append({
                    'id': journal.id,
                    'title': journal.filename,
                    'responsible': journal.editor_in_charge,
                    'remarks': journal.remarks,
                    'stages': {
                        'fee': journal.page_fee,
                        'proof': journal.proof_status
                    },
                    'editors': {
                        'proof12': journal.first_second_proof_editor,
                        'proof3': journal.third_proof_editor,
                        'proofFinal': journal.final_proof_editor
                    },
                    'proofDates': {
                        'editor': journal.editor_time or '',
                        'proof': journal.proof_time or ''
                    }
                })
            
            return JsonResponse({
                'status': 'success',
                'data': new_records
            })
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=500)
    return JsonResponse({'status': 'error', 'message': '不支持的请求方法'}, status=405)

@csrf_exempt
def append_journal_files(request, period_number):
    if request.method == 'POST':
        try:
            # 根据期数选择对应的模型
            model_map = {
                1: JournalPeriodOne,
                2: JournalPeriodTwo,
                3: JournalPeriodThree,
                4: JournalPeriodFour,
                5: JournalPeriodFive
            }
            
            JournalModel = model_map.get(period_number)
            if not JournalModel:
                return JsonResponse({'status': 'error', 'message': '无效的期数'}, status=400)

            new_records = []

            # 优先按文件名数组创建（避免上传内容）
            try:
                if not request.FILES:
                    payload = json.loads(request.body or '{}')
                    filenames = payload.get('filenames', [])
                    if isinstance(filenames, list) and filenames:
                        for name in filenames:
                            journal = JournalModel.objects.create(
                                filename=str(name),
                                editor_in_charge='',
                                page_fee='待更新',
                                proof_status='待更新',
                                first_second_proof_editor='',
                                third_proof_editor='',
                                final_proof_editor='',
                                remarks=''
                            )
                            new_records.append({
                                'id': journal.id,
                                'title': journal.filename,
                                'responsible': journal.editor_in_charge,
                                'remarks': journal.remarks,
                                'stages': {
                                    'fee': journal.page_fee,
                                    'proof': journal.proof_status
                                },
                                'editors': {
                                    'proof12': journal.first_second_proof_editor,
                                    'proof3': journal.third_proof_editor,
                                    'proofFinal': journal.final_proof_editor
                                },
                                'proofDates': {
                                    'editor': journal.editor_time or '',
                                    'proof': journal.proof_time or ''
                                }
                            })
            except Exception:
                pass

            # 兼容旧逻辑：如有上传内容则按文件创建
            if not new_records:
                files = request.FILES.getlist('files')
                for file in files:
                    journal = JournalModel.objects.create(
                        filename=file.name,
                        editor_in_charge='',
                        page_fee='待更新',
                        proof_status='待更新',
                        first_second_proof_editor='',
                        third_proof_editor='',
                        final_proof_editor='',
                        remarks=''
                    )
                    new_records.append({
                        'id': journal.id,
                        'title': journal.filename,
                        'responsible': journal.editor_in_charge,
                        'remarks': journal.remarks,
                        'stages': {
                            'fee': journal.page_fee,
                            'proof': journal.proof_status
                        },
                        'editors': {
                            'proof12': journal.first_second_proof_editor,
                            'proof3': journal.third_proof_editor,
                            'proofFinal': journal.final_proof_editor
                        },
                        'proofDates': {
                            'editor': journal.editor_time or '',
                            'proof': journal.proof_time or ''
                        }
                    })
            
            return JsonResponse({
                'status': 'success',
                'data': new_records
            })
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=500)

@csrf_exempt
def update_journal_field(request, period_number):
    if request.method == 'POST':
        try:
            # 根据期数选择对应的模型
            model_map = {
                1: JournalPeriodOne,
                2: JournalPeriodTwo,
                3: JournalPeriodThree,
                4: JournalPeriodFour,
                5: JournalPeriodFive
            }
            
            JournalModel = model_map.get(period_number)
            if not JournalModel:
                return JsonResponse({'status': 'error', 'message': '无效的期数'}, status=400)

            data = json.loads(request.body)
            journal_id = data.get('id')
            field = data.get('field')
            value = data.get('value')
            
            if not all([journal_id, field]):
                return JsonResponse({'status': 'error', 'message': '缺少必要参数'}, status=400)
            
            journal = JournalModel.objects.get(id=journal_id)
            
            field_mapping = {
                'title': 'filename',
                'responsible': 'editor_in_charge',
                'stages.fee': 'page_fee',
                'stages.proof': 'proof_status',
                'editors.proof12': 'first_second_proof_editor',
                'editors.proof3': 'third_proof_editor',
                'editors.proofFinal': 'final_proof_editor',
                'proofDates.editor': 'editor_time',
                'proofDates.proof': 'proof_time',
                'remarks': 'remarks'
            }
            
            if field in field_mapping:
                if field.startswith('proofDates.'):
                    # 处理时间字段更新
                    field_name = field.split('.')[1]
                    if field_name == 'editor':
                        journal.editor_time = value
                    elif field_name == 'proof':
                        journal.proof_time = value
                elif field.startswith('editors.'):
                    # 处理编辑者字段更新
                    field_name = field.split('.')[1]
                    if field_name == 'proof12':
                        journal.first_second_proof_editor = value
                    elif field_name == 'proof3':
                        journal.third_proof_editor = value
                    elif field_name == 'proofFinal':
                        journal.final_proof_editor = value
                elif field == 'remarks':
                    journal.remarks = value
                else:
                    setattr(journal, field_mapping[field], value)
                journal.save()
                return JsonResponse({'status': 'success'})
            else:
                return JsonResponse({'status': 'error', 'message': f'未知字段: {field}'}, status=400)
                
        except JournalModel.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': f'找不到ID为 {journal_id} 的记录'}, status=404)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def create_journal(request, period_number):
    if request.method == 'POST':
        try:
            # 根据期数选择对应的模型
            model_map = {
                1: JournalPeriodOne,
                2: JournalPeriodTwo,
                3: JournalPeriodThree,
                4: JournalPeriodFour,
                5: JournalPeriodFive
            }
            
            JournalModel = model_map.get(period_number)
            if not JournalModel:
                return JsonResponse({'status': 'error', 'message': '无效的期数'}, status=400)

            data = json.loads(request.body)
            journal = JournalModel.objects.create(
                filename=data.get('title', ''),
                editor_in_charge=data.get('responsible', ''),
                page_fee='待更新',
                proof_status='待更新',
                remarks=data.get('remarks', '')
            )
            
            return JsonResponse({
                'status': 'success',
                'id': journal.id
            })
            
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def clear_journal_data(request, period_number):
    if request.method == 'POST':
        try:
            # 根据期数选择对应的模型
            model_map = {
                1: JournalPeriodOne,
                2: JournalPeriodTwo,
                3: JournalPeriodThree,
                4: JournalPeriodFour,
                5: JournalPeriodFive
            }
            
            JournalModel = model_map.get(period_number)
            if not JournalModel:
                return JsonResponse({'status': 'error', 'message': '无效的期数'}, status=400)

            JournalModel.objects.all().delete()
            return JsonResponse({'status': 'success'})
            
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def delete_journal(request, period_number):
    if request.method == 'POST':
        try:
            # 根据期数选择对应的模型
            model_map = {
                1: JournalPeriodOne,
                2: JournalPeriodTwo,
                3: JournalPeriodThree,
                4: JournalPeriodFour,
                5: JournalPeriodFive
            }
            
            JournalModel = model_map.get(period_number)
            if not JournalModel:
                return JsonResponse({'status': 'error', 'message': '无效的期数'}, status=400)

            data = json.loads(request.body)
            journal_id = data.get('id')
            
            if not journal_id:
                return JsonResponse({'status': 'error', 'message': '缺少记录ID'}, status=400)
            
            journal = JournalModel.objects.get(id=journal_id)
            journal.delete()
            
            return JsonResponse({'status': 'success'})
            
        except JournalModel.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': f'找不到ID为 {journal_id} 的记录'}, status=404)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': '不支持的请求方法'}, status=405)

@csrf_exempt
def get_period_mapping(request):
    try:
        mappings = PeriodMapping.objects.all()
        data = {
            mapping.backend_period: mapping.display_period 
            for mapping in mappings
        }
        return JsonResponse(data)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def update_period_mapping(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            backend_period = data.get('backend_period')
            display_period = data.get('display_period')
            
            if not all([backend_period, display_period]):
                return JsonResponse({'status': 'error', 'message': '缺少必要参数'}, status=400)
                
            mapping, created = PeriodMapping.objects.update_or_create(
                backend_period=backend_period,
                defaults={'display_period': display_period}
            )
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def load_template(request):
    try:
        template_name = request.GET.get('template')
        if not template_name:
            return JsonResponse({'status': 'error', 'message': '模板名称不能为空'})
            
        # 构建模板文件路径
        template_path = os.path.join(settings.BASE_DIR, 'qikan', 'templates', template_name)
        
        # 检查文件是否存在
        if not os.path.exists(template_path):
            # 如果文件不存在，创建一个空的模板文件
            doc = Document()
            doc.add_paragraph('这是一个新的模板文件')
            doc.save(template_path)
            return JsonResponse({
                'status': 'success',
                'content': '这是一个新的模板文件'
            })

        # 读取 Word 文档
        doc = Document(template_path)
        
        # 提取文本内容
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
            
        return JsonResponse({
            'status': 'success',
            'content': '\n'.join(content)
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@csrf_exempt
def save_template(request):
    try:
        data = json.loads(request.body)
        template_name = data.get('path')
        content = data.get('content')
        
        if not all([template_name, content]):
            return JsonResponse({'status': 'error', 'message': '模板名称和内容不能为空'})
            
        # 构建模板文件路径
        template_path = os.path.join(settings.BASE_DIR, 'qikan', 'templates', template_name)
        
        # 创建新的 Word 文档
        doc = Document()
        
        # 将内容按行分割并添加到文档中
        for line in content.split('\n'):
            doc.add_paragraph(line)
            
        # 保存文档
        doc.save(template_path)
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@csrf_exempt
def generate_notification(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            params_list = data.get('params_list', [])
            task_type = data.get('task_type', 'A')  # 'L'=录用通知, 'B'=版面费通知, 'A'=全部
            
            # 确定输出目录
            output_folder = os.path.join(settings.BASE_DIR, 'qikan', 'output_pdfs')
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
            
            generated_files = []
            
            # 执行录用通知任务
            if task_type in ('L', 'A'):
                acceptance_template = os.path.join(settings.BASE_DIR, 'qikan', 'templates', 'accept_template.docx')
                if os.path.exists(acceptance_template):
                    generate_notice_pdfs(
                        data_list=params_list,
                        template_path=acceptance_template,
                        output_prefix='录用通知',
                        output_folder=output_folder,
                        delete_docx_after_success=True
                    )
                else:
                    return JsonResponse({
                        'status': 'error',
                        'message': '录用通知模板文件未找到'
                    }, status=404)
            
            # 执行版面费通知任务
            if task_type in ('B', 'A'):
                charge_template = os.path.join(settings.BASE_DIR, 'qikan', 'templates', 'charge_template.docx')
                if os.path.exists(charge_template):
                    generate_notice_pdfs(
                        data_list=params_list,
                        template_path=charge_template,
                        output_prefix='版面费通知',
                        output_folder=output_folder,
                        delete_docx_after_success=True
                    )
                else:
                    return JsonResponse({
                        'status': 'error',
                        'message': '版面费通知模板文件未找到'
                    }, status=404)
            
            # 获取生成的文件列表
            for filename in os.listdir(output_folder):
                if filename.endswith('.pdf'):
                    file_path = os.path.join(output_folder, filename)
                    file_size = os.path.getsize(file_path)
                    generated_files.append({
                        'name': filename,
                        'size': file_size,
                        'path': os.path.join('output_pdfs', filename)
                    })
            
            if not generated_files:
                return JsonResponse({
                    'status': 'error',
                    'message': '未生成任何文件'
                }, status=400)
            
            return JsonResponse({
                'status': 'success',
                'message': f'成功生成 {len(generated_files)} 个文件',
                'files': generated_files
            })
            
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=500)

@csrf_exempt
def download_notification(request, filename):
    try:
        file_path = os.path.join(settings.BASE_DIR, 'qikan', 'output_pdfs', filename)
        if os.path.exists(file_path):
            response = FileResponse(open(file_path, 'rb'))
            response['Content-Type'] = 'application/pdf'
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        else:
            return JsonResponse({
                'status': 'error',
                'message': '文件不存在'
            }, status=404)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@csrf_exempt
def clear_notification_files(request):
    """清理output_pdfs文件夹中的所有文件"""
    if request.method == 'POST':
        try:
            output_folder = os.path.join(settings.BASE_DIR, 'qikan', 'output_pdfs')
            if os.path.exists(output_folder):
                for filename in os.listdir(output_folder):
                    file_path = os.path.join(output_folder, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.unlink(file_path)
                    except Exception as e:
                        print(f"Error: {e}")
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=500)

class ContactViewSet(viewsets.ModelViewSet):
    queryset = Contact.objects.all()
    serializer_class = ContactSerializer

    @action(detail=False, methods=['post'])
    def batch_create(self, request):
        """批量导入联系人数据"""
        serializer = ContactSerializer(data=request.data.get('contacts', []), many=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['delete'])
    def batch_delete(self, request):
        """批量删除联系人数据"""
        label = request.data.get('label')
        if label:
            Contact.objects.filter(label=label).delete()
        else:
            Contact.objects.all().delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

class ReviewerViewSet(viewsets.ModelViewSet):
    """责编信息视图集"""
    queryset = Reviewer.objects.all()
    serializer_class = ReviewerSerializer

    @action(detail=False, methods=['post'])
    def batch_create(self, request):
        """批量创建责编信息"""
        reviewers_data = request.data.get('reviewers', [])
        serializer = self.get_serializer(data=reviewers_data, many=True)
        
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['delete'])
    def batch_delete(self, request):
        """批量删除责编信息"""
        try:
            self.get_queryset().delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

# 待办事项相关视图
@csrf_exempt
@require_http_methods(["GET", "POST"])
def todo_list_view(request):
    """获取待办事项列表或创建新的待办事项"""
    if request.method == 'GET':
        try:
            # 获取筛选参数
            status_filter = request.GET.get('status', '')
            assignee_filter = request.GET.get('assignee', '')
            priority_filter = request.GET.get('priority', '')
            
            # 构建查询
            queryset = Todo.objects.all()
            
            if status_filter:
                queryset = queryset.filter(status=status_filter)
            if assignee_filter:
                queryset = queryset.filter(assignee=assignee_filter)
            if priority_filter:
                queryset = queryset.filter(priority=priority_filter)
            
            # 排序：优先级(高->中->低) -> 状态(进行中->已完成) -> 创建时间(新->旧)
            priority_order = {'high': 3, 'medium': 2, 'low': 1}
            status_order = {'in_progress': 2, 'completed': 1}
            
            todos = []
            for todo in queryset:
                todos.append({
                    'id': todo.id,
                    'content': todo.content,
                    'assignee': todo.assignee,
                    'priority': todo.priority,
                    'status': todo.status,
                    'dueDate': todo.due_date.isoformat() if todo.due_date else None,
                    'createdAt': todo.created_at.isoformat(),
                    'updatedAt': todo.updated_at.isoformat(),
                })
            
            # Python排序
            todos.sort(key=lambda x: (
                -priority_order.get(x['priority'], 0),
                -status_order.get(x['status'], 0),
                -int(x['createdAt'].replace('-', '').replace(':', '').replace('T', '').replace('.', '')[:14])
            ))
            
            return JsonResponse({
                'success': True,
                'data': todos
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)
    
    elif request.method == 'POST':
        try:
            data = json.loads(request.body)
            
            # 验证必填字段
            if not data.get('content', '').strip():
                return JsonResponse({
                    'success': False,
                    'error': '任务内容不能为空'
                }, status=400)
            
            if not data.get('assignee'):
                return JsonResponse({
                    'success': False,
                    'error': '请选择添加人'
                }, status=400)
            
            # 创建待办事项
            todo = Todo.objects.create(
                content=data['content'].strip(),
                assignee=data['assignee'],
                priority=data.get('priority', 'medium'),
                status=data.get('status', 'in_progress'),
                due_date=parse_date(data['dueDate']) if data.get('dueDate') else None
            )
            
            return JsonResponse({
                'success': True,
                'data': {
                    'id': todo.id,
                    'content': todo.content,
                    'assignee': todo.assignee,
                    'priority': todo.priority,
                    'status': todo.status,
                    'dueDate': todo.due_date.isoformat() if todo.due_date else None,
                    'createdAt': todo.created_at.isoformat(),
                    'updatedAt': todo.updated_at.isoformat(),
                }
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)

@csrf_exempt
@require_http_methods(["PUT", "DELETE"])
def todo_detail_view(request, todo_id):
    """更新或删除特定的待办事项"""
    try:
        todo = Todo.objects.get(id=todo_id)
    except Todo.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': '待办事项不存在'
        }, status=404)
    
    if request.method == 'PUT':
        try:
            data = json.loads(request.body)
            
            # 验证必填字段
            if not data.get('content', '').strip():
                return JsonResponse({
                    'success': False,
                    'error': '任务内容不能为空'
                }, status=400)
            
            # 更新字段
            todo.content = data['content'].strip()
            todo.assignee = data.get('assignee', todo.assignee)
            todo.priority = data.get('priority', todo.priority)
            todo.status = data.get('status', todo.status)
            todo.due_date = parse_date(data['dueDate']) if data.get('dueDate') else None
            todo.save()
            
            return JsonResponse({
                'success': True,
                'data': {
                    'id': todo.id,
                    'content': todo.content,
                    'assignee': todo.assignee,
                    'priority': todo.priority,
                    'status': todo.status,
                    'dueDate': todo.due_date.isoformat() if todo.due_date else None,
                    'createdAt': todo.created_at.isoformat(),
                    'updatedAt': todo.updated_at.isoformat(),
                }
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)
    
    elif request.method == 'DELETE':
        try:
            todo.delete()
            return JsonResponse({
                'success': True,
                'message': '待办事项已删除'
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)

@csrf_exempt
@require_http_methods(["PATCH"])
def todo_status_view(request, todo_id):
    """更新待办事项状态"""
    try:
        todo = Todo.objects.get(id=todo_id)
    except Todo.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': '待办事项不存在'
        }, status=404)
    
    try:
        data = json.loads(request.body)
        new_status = data.get('status')
        
        if new_status not in ['in_progress', 'completed']:
            return JsonResponse({
                'success': False,
                'error': '无效的状态值'
            }, status=400)
        
        todo.status = new_status
        todo.save()
        
        return JsonResponse({
            'success': True,
            'data': {
                'id': todo.id,
                'status': todo.status,
                'updatedAt': todo.updated_at.isoformat(),
            }
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)

@csrf_exempt
@require_http_methods(["POST"])
def check_references(request):
    """参考文献纠错API"""
    try:
        from .RC import ReferenceChecker
        
        data = json.loads(request.body)
        references_text = data.get('references', '').strip()
        
        if not references_text:
            return JsonResponse({
                'status': 'error',
                'message': '请输入参考文献内容'
            }, status=400)
        
        # 创建检查器实例
        checker = ReferenceChecker()
        
        # 将字符串按行分割，并去除空行
        reference_list = [line.strip() for line in references_text.strip().split('\n') if line.strip()]
        
        if not reference_list:
            return JsonResponse({
                'status': 'error',
                'message': '未检测到有效的参考文献条目'
            }, status=400)
        
        # 构建结果字符串
        result_lines = []
        result_lines.append("--- 以下为自动化程序检查结果，最终以\"中文信息学报参考文献规范.docx\"文件为准 ---")
        result_lines.append("")
        result_lines.append("")
        
        # 遍历列表，逐条检查并输出结果
        for ref in reference_list:
            if not ref:  # 跳过可能的空行
                continue
            
            result_lines.append(f"{ref}")
            errors = checker.check_reference(ref)
            for error in errors:
                result_lines.append(f"  - 问题: {error}")
            result_lines.append("")  # 每条参考文献后添加空行
        
        result_text = '\n'.join(result_lines)
        
        return JsonResponse({
            'status': 'success',
            'result': result_text
        })
        
    except ImportError:
        return JsonResponse({
            'status': 'error',
            'message': '参考文献检查模块未找到，请确保RC.py文件存在'
        }, status=500)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'检查过程中出现错误: {str(e)}'
        }, status=500)


# 邮件模板管理API
from .models import EmailTemplate
from django.core.paginator import Paginator

@csrf_exempt
@require_http_methods(["GET"])
def get_email_templates(request):
    """获取邮件模板列表"""
    print(f"接收到获取模板列表请求: {request.method}")
    try:
        templates = EmailTemplate.objects.all().order_by('-created_at')
        template_list = []
        
        for template in templates:
            template_list.append({
                'id': template.id,
                'title': template.title,
                'content': template.content,
                'createTime': template.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updateTime': template.updated_at.strftime('%Y-%m-%d %H:%M:%S')
            })
        
        print(f"返回模板数量: {len(template_list)}")
        resp = JsonResponse({
            'status': 'success',
            'data': template_list
        })
        # 避免浏览器/代理缓存模板列表，确保前端获取最新内容
        resp['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        resp['Pragma'] = 'no-cache'
        resp['Expires'] = '0'
        return resp
    except Exception as e:
        print(f"获取模板列表错误: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'message': f'获取模板列表失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def create_email_template(request):
    """创建邮件模板"""
    print(f"接收到创建模板请求: {request.method}")
    print(f"请求体: {request.body}")
    try:
        data = json.loads(request.body)
        title = data.get('title', '').strip()
        content = data.get('content', '').strip()
        
        print(f"解析数据 - 标题: {title}, 内容长度: {len(content)}")
        
        if not title:
            return JsonResponse({
                'status': 'error',
                'message': '模板标题不能为空'
            }, status=400)
        
        template = EmailTemplate.objects.create(
            title=title,
            content=content
        )
        print(f"模板创建成功，ID: {template.id}")
        
        return JsonResponse({
            'status': 'success',
            'data': {
                'id': template.id,
                'title': template.title,
                'content': template.content,
                'createTime': template.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updateTime': template.updated_at.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'创建模板失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["PUT"])
def update_email_template(request, template_id):
    """更新邮件模板"""
    try:
        template = EmailTemplate.objects.get(id=template_id)
        data = json.loads(request.body)
        
        title = data.get('title', '').strip()
        content = data.get('content', '').strip()
        
        if not title:
            return JsonResponse({
                'status': 'error',
                'message': '模板标题不能为空'
            }, status=400)
        
        template.title = title
        template.content = content
        template.save()
        
        return JsonResponse({
            'status': 'success',
            'data': {
                'id': template.id,
                'title': template.title,
                'content': template.content,
                'createTime': template.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'updateTime': template.updated_at.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
    except EmailTemplate.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': '模板不存在'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'更新模板失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["DELETE"])
def delete_email_template(request, template_id):
    """删除邮件模板"""
    try:
        template = EmailTemplate.objects.get(id=template_id)
        template.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': '模板删除成功'
        })
    except EmailTemplate.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': '模板不存在'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'删除模板失败: {str(e)}'
        }, status=500)


# ==================== 联系人管理 API ====================

@csrf_exempt
@require_http_methods(["GET"])
def get_recipients(request):
    """获取联系人列表"""
    try:
        search = request.GET.get('search', '')
        group_name = request.GET.get('group_name', '')
        page = int(request.GET.get('page', 1))
        page_size = int(request.GET.get('page_size', 10))
        
        recipients = Recipient.objects.all().prefetch_related('groups')
        
        # 搜索过滤
        if search:
            recipients = recipients.filter(
                Q(name__icontains=search) | Q(email__icontains=search)
            )
            # 服务器端限制搜索时的最大返回量（Top 50），用于低带宽环境
            if page_size > 50:
                page_size = 50
        
        # 分组过滤
        if group_name:
            # 前端传入的 group_name 一律按真实分组名称过滤
            recipients = recipients.filter(groups__name=group_name)
        
        total_count = recipients.count()
        
        # 分页
        start = (page - 1) * page_size
        end = start + page_size
        recipients = recipients[start:end]
        
        # 获取所有分组及其统计
        # 仅返回真实存在的分组及其成员数
        group_list = [{'name': g.name, 'count': g.members.count()} for g in RecipientGroup.objects.all()]
        
        recipient_list = []
        for recipient in recipients:
            # 兼容：返回单个 groupName（若多组则取第一个），前端已有多处依赖
            group_names = [g.name for g in recipient.groups.all()]
            group_name_val = group_names[0] if group_names else '默认分组'
            recipient_list.append({
                'id': recipient.id,
                'name': recipient.name,
                'email': recipient.email,
                'groupName': group_name_val,
                'groups': group_names,
                'remark': recipient.remark or '',
                'createTime': recipient.created_at.strftime('%Y-%m-%d %H:%M:%S')
            })
        
        resp_body = {
            'status': 'success',
            'data': {
                'recipients': recipient_list,
                'groups': group_list,
                'total': total_count,
                'page': page,
                'pageSize': page_size,
                'totalPages': (total_count + page_size - 1) // page_size
            }
        }
        # 计算 ETag 并处理条件请求
        payload = json.dumps(resp_body, ensure_ascii=False, separators=(",", ":")).encode('utf-8')
        etag = 'W/"' + hashlib.md5(payload).hexdigest() + '"'
        inm = request.headers.get('If-None-Match') or request.META.get('HTTP_IF_NONE_MATCH')
        if inm and inm == etag:
            return HttpResponseNotModified()
        resp = JsonResponse(resp_body)
        resp['ETag'] = etag
        return resp
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'获取联系人列表失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def create_recipient(request):
    """创建联系人（允许同邮箱重复，始终新增并按需追加分组）"""
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        email = data.get('email', '').strip()
        group_name = data.get('groupName', '').strip()
        remark = data.get('remark', '').strip()

        if not email:
            return JsonResponse({'status': 'error', 'message': '邮箱不能为空'}, status=400)
        if not name:
            name = email.split('@')[0]

        recipient = Recipient.objects.create(name=name, email=email, remark=remark)
        if group_name and group_name != '默认分组':
            group_obj, _ = RecipientGroup.objects.get_or_create(name=group_name)
            recipient.groups.add(group_obj)

        return JsonResponse({
            'status': 'success',
            'message': '联系人创建成功',
            'data': {
                'id': recipient.id,
                'name': recipient.name,
                'email': recipient.email,
                'groupName': group_name or '默认分组',
                'remark': recipient.remark,
                'createTime': recipient.created_at.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'创建联系人失败: {str(e)}'}, status=500)

@csrf_exempt
@require_http_methods(["PUT"])
def update_recipient(request, recipient_id):
    """更新联系人"""
    try:
        data = json.loads(request.body)
        recipient = Recipient.objects.get(id=recipient_id)
        
        name = data.get('name', '').strip()
        email = data.get('email', '').strip()
        group_name = data.get('groupName', '').strip()
        remark = data.get('remark', '').strip()
        
        if not name or not email:
            return JsonResponse({
                'status': 'error',
                'message': '姓名和邮箱不能为空'
            }, status=400)
        
        # 检查邮箱是否已被其他联系人使用
        if Recipient.objects.filter(email=email).exclude(id=recipient_id).exists():
            return JsonResponse({
                'status': 'error',
                'message': '该邮箱已被其他联系人使用'
            }, status=400)
        
        recipient.name = name
        recipient.email = email
        recipient.remark = remark
        recipient.save()
        # 兼容旧行为：设置为“移动到某组”，因此先清空再设定单一分组
        recipient.groups.clear()
        if group_name and group_name != '默认分组':
            group_obj, _ = RecipientGroup.objects.get_or_create(name=group_name)
            recipient.groups.add(group_obj)
        
        return JsonResponse({
            'status': 'success',
            'message': '联系人更新成功',
            'data': {
                'id': recipient.id,
                'name': recipient.name,
                'email': recipient.email,
                'groupName': group_name or '默认分组',
                'remark': recipient.remark,
                'createTime': recipient.created_at.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
    except Recipient.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': '联系人不存在'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'更新联系人失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["DELETE"])
def delete_recipient(request, recipient_id):
    """删除联系人"""
    try:
        recipient = Recipient.objects.get(id=recipient_id)
        recipient.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': '联系人删除成功'
        })
    except Recipient.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': '联系人不存在'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'删除联系人失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def batch_delete_recipients(request):
    """批量删除联系人"""
    try:
        data = json.loads(request.body)
        recipient_ids = data.get('ids', [])
        
        if not recipient_ids:
            return JsonResponse({
                'status': 'error',
                'message': '请选择要删除的联系人'
            }, status=400)
        
        deleted_count = Recipient.objects.filter(id__in=recipient_ids).delete()[0]
        
        return JsonResponse({
            'status': 'success',
            'message': f'成功删除 {deleted_count} 个联系人'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'批量删除失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def batch_update_group(request):
    """批量更新联系人分组"""
    try:
        data = json.loads(request.body)
        recipient_ids = data.get('ids', [])
        group_name = data.get('groupName', '').strip()
        
        if not recipient_ids:
            return JsonResponse({
                'status': 'error',
                'message': '请选择要更新的联系人'
            }, status=400)
        
        updated_count = 0
        if group_name and group_name != '默认分组':
            group_obj, _ = RecipientGroup.objects.get_or_create(name=group_name)
            for r in Recipient.objects.filter(id__in=recipient_ids):
                r.groups.clear()
                r.groups.add(group_obj)
                updated_count += 1
        else:
            for r in Recipient.objects.filter(id__in=recipient_ids):
                r.groups.clear()
                updated_count += 1
        
        return JsonResponse({
            'status': 'success',
            'message': f'成功更新 {updated_count} 个联系人的分组'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'批量更新分组失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def import_recipients(request):
    """批量导入联系人（允许同邮箱重复，始终新增并按需追加分组）。
    规则：
    - 不以邮箱为唯一键，任何行都可创建新联系人；
    - 若提供了 groupName 且不为“默认分组”，则为该联系人追加该分组；
    - 仅对缺失邮箱的行计为错误；姓名缺失用邮箱前缀兜底；
    返回：createdCount、updatedCount(恒为0)、errorCount、errors。
    """
    try:
        data = json.loads(request.body)
        recipients_data = data.get('recipients', [])
        group_name = data.get('groupName', '').strip()

        if not recipients_data:
            return JsonResponse({
                'status': 'error',
                'message': '没有要导入的联系人数据'
            }, status=400)

        created_count = 0
        updated_count = 0  # 兼容旧字段
        error_list = []

        group_obj = None
        if group_name and group_name != '默认分组':
            group_obj, _ = RecipientGroup.objects.get_or_create(name=group_name)

        for i, recipient_data in enumerate(recipients_data):
            try:
                name = (recipient_data.get('name') or '').strip()
                email = (recipient_data.get('email') or '').strip()
                remark = (recipient_data.get('remark') or '').strip()

                if not email:
                    error_list.append(f'第{i+1}行: 邮箱不能为空')
                    continue
                # 若缺少姓名，允许用邮箱前缀兜底
                if not name:
                    name = email.split('@')[0]

                r = Recipient.objects.create(name=name, email=email, remark=remark)
                if group_obj:
                    r.groups.add(group_obj)
                created_count += 1

            except Exception as e:
                error_list.append(f'第{i+1}行: {str(e)}')

        return JsonResponse({
            'status': 'success',
            'message': f'新增 {created_count} 个',
            'data': {
                'createdCount': created_count,
                'updatedCount': updated_count,
                'errorCount': len(error_list),
                'errors': error_list
            }
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'导入联系人失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["GET"])
def get_groups(request):
    """获取所有分组及统计"""
    try:
        # 统计分组人数（membership）
        group_list = []
        for g in RecipientGroup.objects.all():
            group_list.append({'name': g.name, 'count': g.members.count()})
        # 计算 ETag（基于序列化后的内容）
        payload = json.dumps({'status': 'success', 'data': group_list}, ensure_ascii=False, separators=(",", ":")).encode('utf-8')
        etag = 'W/"' + hashlib.md5(payload).hexdigest() + '"'

        # 条件请求：If-None-Match
        inm = request.headers.get('If-None-Match') or request.META.get('HTTP_IF_NONE_MATCH')
        if inm and inm == etag:
            return HttpResponseNotModified()

        resp = JsonResponse({'status': 'success', 'data': group_list})
        resp['ETag'] = etag
        return resp
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'获取分组列表失败: {str(e)}'
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def append_to_group(request):
    """在不清除原分组的前提下，将收件人追加到指定分组（支持批量）。
    请求体: { ids: number[] | number, groupName: string }
    """
    try:
        data = json.loads(request.body)
        ids = data.get('ids')
        if isinstance(ids, int):
            ids = [ids]
        if not ids or not isinstance(ids, list):
            return JsonResponse({'status': 'error', 'message': 'ids 必须为非空数组或数字'}, status=400)
        group_name = (data.get('groupName') or '').strip()
        if not group_name:
            return JsonResponse({'status': 'error', 'message': 'groupName 不能为空'}, status=400)

        try:
            group = RecipientGroup.objects.get(name=group_name)
        except RecipientGroup.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': '分组不存在'}, status=404)

        affected = 0
        for rid in ids:
            try:
                r = Recipient.objects.get(id=rid)
            except Recipient.DoesNotExist:
                continue
            r.groups.add(group)  # 追加，不清除已有分组
            # 若追加到的不是“默认分组”，则移除其在真实“默认分组”中的成员关系（若存在）
            if group.name != '默认分组':
                try:
                    default_group = RecipientGroup.objects.get(name='默认分组')
                    r.groups.remove(default_group)
                except RecipientGroup.DoesNotExist:
                    pass
            affected += 1

        return JsonResponse({'status': 'success', 'data': {'affected': affected}})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'追加分组失败: {str(e)}'}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def remove_from_group(request):
    """从指定分组移除收件人（支持批量），不影响该收件人的其他分组。
    请求体: { ids: number[] | number, groupName: string }
    """
    try:
        data = json.loads(request.body)
        ids = data.get('ids')
        if isinstance(ids, int):
            ids = [ids]
        if not ids or not isinstance(ids, list):
            return JsonResponse({'status': 'error', 'message': 'ids 必须为非空数组或数字'}, status=400)
        group_name = (data.get('groupName') or '').strip()
        if not group_name:
            return JsonResponse({'status': 'error', 'message': 'groupName 不能为空'}, status=400)

        try:
            group = RecipientGroup.objects.get(name=group_name)
        except RecipientGroup.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': '分组不存在'}, status=404)

        affected = 0
        for rid in ids:
            try:
                r = Recipient.objects.get(id=rid)
            except Recipient.DoesNotExist:
                continue
            r.groups.remove(group)
            # 若移出后没有任何分组，则自动加入真实“默认分组”（若存在）
            if r.groups.count() == 0:
                try:
                    default_group = RecipientGroup.objects.get(name='默认分组')
                    r.groups.add(default_group)
                except RecipientGroup.DoesNotExist:
                    pass
            affected += 1

        return JsonResponse({'status': 'success', 'data': {'affected': affected}})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'移出分组失败: {str(e)}'}, status=500)


@csrf_exempt
@require_http_methods(["POST"])  # 管理端临时使用：修复“默认分组”唯一性约束
def cleanup_ungrouped_invariant(request):
    """一次性修复数据，使“默认分组”只在无其它分组时存在：
    - 对同时存在“默认分组”和其它分组的联系人：移除“默认分组”。
    - 对没有任何分组的联系人：若存在真实“默认分组”，则加入其中。
    返回：修复的计数。
    """
    try:
        try:
            default_group = RecipientGroup.objects.get(name='默认分组')
        except RecipientGroup.DoesNotExist:
            default_group = None

        removed_default = 0
        added_default = 0

        for r in Recipient.objects.all().prefetch_related('groups'):
            gs = list(r.groups.all())
            if default_group and any(g.id == default_group.id for g in gs) and len(gs) > 1:
                # 同时存在“默认分组”和其它分组：移除“默认分组”
                r.groups.remove(default_group)
                removed_default += 1
            elif len(gs) == 0 and default_group:
                # 没有任何分组：加入“默认分组”
                r.groups.add(default_group)
                added_default += 1

        return JsonResponse({'status': 'success', 'data': {
            'removed_default': removed_default,
            'added_default': added_default,
        }})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'清理失败: {str(e)}'}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def create_group(request):
    """新建分组（用于创建空分组或预创建分组）"""
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        if not name:
            return JsonResponse({'status': 'error', 'message': '分组名称不能为空'}, status=400)
        if name == '默认分组':
            return JsonResponse({'status': 'error', 'message': '该名称保留为默认分组'}, status=400)
        group, created = RecipientGroup.objects.get_or_create(name=name)
        return JsonResponse({'status': 'success', 'data': {'name': group.name, 'created': created}})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'创建分组失败: {str(e)}'}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def rename_group(request):
    """重命名分组"""
    try:
        data = json.loads(request.body)
        old_name = data.get('oldName', '').strip()
        new_name = data.get('newName', '').strip()
        
        if not old_name or not new_name:
            return JsonResponse({
                'status': 'error',
                'message': '分组名称不能为空'
            }, status=400)
        
        if old_name == new_name:
            return JsonResponse({
                'status': 'error',
                'message': '新分组名称与原名称相同'
            }, status=400)
        
        if old_name == '默认分组' or new_name == '默认分组':
            return JsonResponse({
                'status': 'error',
                'message': '不能对默认分组执行该操作'
            }, status=400)
        
        # 若新名已存在：合并（把旧组成员移到新组，再删除旧组）
        try:
            new_group = RecipientGroup.objects.get(name=new_name)
            # 合并
            try:
                old_group = RecipientGroup.objects.get(name=old_name)
            except RecipientGroup.DoesNotExist:
                return JsonResponse({'status': 'error', 'message': '原分组不存在'}, status=404)
            moved = 0
            for r in old_group.members.all():
                r.groups.add(new_group)
                moved += 1
            old_group.delete()
            updated_count = moved
        except RecipientGroup.DoesNotExist:
            # 直接改名
            try:
                old_group = RecipientGroup.objects.get(name=old_name)
            except RecipientGroup.DoesNotExist:
                return JsonResponse({'status': 'error', 'message': '原分组不存在'}, status=404)
            old_group.name = new_name
            old_group.save(update_fields=['name'])
            updated_count = old_group.members.count()
        
        return JsonResponse({
            'status': 'success',
            'message': f'成功重命名分组，更新了 {updated_count} 个联系人'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'重命名分组失败: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["DELETE"])
def delete_group(request):
    """删除分组，并根据邮箱唯一性清理：
    - 若某邮箱在全库仅出现一次：仅移除该联系人的该分组标签；
    - 若某邮箱在全库出现多次：删除该邮箱的多余联系人，仅保留一个实例；
    - 最终删除该分组实体。
    """
    try:
        data = json.loads(request.body)
        group_name = data.get('groupName', '').strip()
        
        if not group_name:
            return JsonResponse({
                'status': 'error',
                'message': '分组名称不能为空'
            }, status=400)
        
        if group_name == '默认分组':
            return JsonResponse({
                'status': 'error',
                'message': '不能删除默认分组'
            }, status=400)
        
        try:
            group = RecipientGroup.objects.get(name=group_name)
        except RecipientGroup.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': '分组不存在'}, status=404)
        # 先收集该组内成员，按邮箱聚合
        members = list(group.members.all())
        email_to_members = {}
        for r in members:
            email_to_members.setdefault(r.email, []).append(r)

        affected = 0
        for email, rs in email_to_members.items():
            total_with_email = Recipient.objects.filter(email=email).count()
            if total_with_email <= 1:
                # 唯一邮箱：仅移除该分组标签
                for r in rs:
                    r.groups.remove(group)
                    affected += 1
            else:
                # 重复邮箱：全局保留一个，删除其余
                # 策略：优先保留当前分组外的一个；若都在该分组则保留第一条
                keep = None
                for r in Recipient.objects.filter(email=email).prefetch_related('groups'):
                    if group not in r.groups.all():
                        keep = r
                        break
                if keep is None:
                    keep = rs[0]
                # 删除其它重复项
                for r in Recipient.objects.filter(email=email).exclude(id=keep.id):
                    r.delete()
                # 确保保留项移除该分组
                keep.groups.remove(group)
                affected += 1

        updated_count = affected
        # 最后删除分组实体
        group.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': f'成功删除分组并清理，影响 {updated_count} 条联系人记录'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'删除分组失败: {str(e)}'
        }, status=500)
